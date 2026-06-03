from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class UnsupportedDocumentTypeError(ValueError):
    pass


class EmptyDocumentTextError(ValueError):
    pass


def extract_text_from_upload(filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {extension}")

    if extension == ".txt":
        text = content.decode("utf-8", errors="ignore")
    elif extension == ".pdf":
        text = extract_text_from_pdf(content)
    else:
        text = extract_text_from_docx(content)

    text = text.strip()
    if not text:
        raise EmptyDocumentTextError("Extracted text is empty")

    return text


def extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    page_texts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(page_texts)


def extract_text_from_docx(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells.append(cell.text)
    return "\n".join(paragraphs + table_cells)
